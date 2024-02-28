import re, os, sys,datetime,pytz
from docx import Document
from docx2pdf import convert

print(f" WELCOME TO PARTIAL DOCX EDITOR ".center(50,"~"))
replacement_code_txt="replacement_code.txt" 
if os.path.isfile(f".\{replacement_code_txt}") != True:
    with open(replacement_code_txt,"w") as file:
        file.write("# <Company_name>, <Job_title> and <Recruiter_name> fields automatically asked by the code. These three replacement codes can be used in template docx.\n#You can add more code without starting with #. Type your replacement code line by line to below .\n#<your_replacement_code>:\n")
print("!!! >>> File type of CV template must be in docx")
template_filename =input("\nDefault : cv template\nEnter the CV Word Template name without file type: ")
if template_filename.title() == "Q":
    print("Quited!")   
    quit()
if template_filename == "":
    template_filename = "cv template.docx"
else:
    template_filename += ".docx"

output_filename = input(f"\nDefault : Cover Letter and Resume\nOutput Filename: ").title()
if output_filename == "":
    output_filename ="Cover Letter and Resume"

print(f"\n!!! Don't forget to edit and match \"{replacement_code_txt}\" with \"{template_filename}\"")

def get_current_time():
    now=datetime.datetime.now(pytz.timezone(("Europe/Dublin")))#Current Time according to Selected Timezone
    now_formatted = now.strftime("%Y%m%d_%H%M")
    return now_formatted

def read_user_codes():
    with open (replacement_code_txt,"r") as file:
        user_codes = file.readlines()
        user_codes = [x.strip() for x in user_codes if x.startswith("#") == False ]
        return user_codes


def get_user_input():
    global user_input_dict    
    user_input_dict = dict()
    user_input_dict["<Company_name>"] = ""
    user_input_dict["<Job_title>"] = ""
    user_input_dict["<Recruiter_name>"] = ""

    for code in read_user_codes():
        user_input_dict[code] = ""
    
    for key in user_input_dict.keys():
        if key == "<Recruiter_name>":
            readinput = input(f"\nDefault : Recruiter\n{key}: ")
            if readinput =="":
                readinput = "Recruiter"
        else:
            readinput = input(f"\n{key}: ")
        if readinput.title() == "Q":
            print("Quited!")
            quit()
        if key in ["<Company_name>","<Job_title>","<Recruiter_name>"]:
            user_input_dict[key] = readinput.title()
        else:
            user_input_dict[key] = readinput

def docx_replace_regex(word_document, regex , replace):
    
    for p in word_document.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in word_document.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


while True:
    get_user_input()
    word_document = Document(template_filename)
    for usercode, original in user_input_dict.items():
        usercode_regex =re.compile(usercode)
        docx_replace_regex(word_document, usercode_regex, original)
    company_name = user_input_dict["<Company_name>"].title()
    job_title = user_input_dict["<Job_title>"].title()
    worddocname  = f"{output_filename} for {job_title} at {company_name}_{get_current_time()}.docx"
    word_document.save(worddocname)
    convert(worddocname, f"{worddocname[:-5]}.pdf")
    print("'q' for quit.")
